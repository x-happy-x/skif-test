from copy import deepcopy

from docx.oxml import OxmlElement
from docx.table import Table
from docx.text.paragraph import Paragraph

from . import STYLES, STEPS, TYPE_CLOSED_QUESTIONS, TYPE_OPENED_QUESTIONS, TYPE_ACCORDANCE_QUESTIONS, OpenedQuestion, ClosedQuestion, AccordanceQuestion

find_types = {
    "equal": lambda a, b: a == b,
    "start": lambda a, b: a.startswith(b),
    "end": lambda a, b: a.endswith(b),
    "contains": lambda a, b: a.find(b) >= 0
}


def insert_paragraph_after(paragraph, text=None, style=None):
    new_paragraph = OxmlElement("w:p")
    paragraph._element.addnext(new_paragraph)
    new_para = Paragraph(new_paragraph, paragraph._parent)
    if text:
        new_para.text = text
    if style:
        new_para.style = style
    return new_para


def find_paragraph_id(template, pattern, find, start_id=0):
    test_paragraph_id = -1
    for i in range(start_id, len(template.paragraphs)):
        if find(template.paragraphs[i].text.strip(), pattern):
            test_paragraph_id = i
            break
    return test_paragraph_id


def move_table_after(table, paragraph):
    tbl, p = table._element, paragraph._element
    p.addnext(tbl)
    return table


def copy_table_after(table, paragraph):
    tbl, p = table._element, paragraph._element
    new_tbl = deepcopy(tbl)
    p.addnext(new_tbl)
    return Table(new_tbl, p._parent)


def replace_test(document, paragraph, test_type, questions, steps=None, styles=None, questions_start=0):
    if steps is None:
        steps = STEPS[test_type]
    if styles is None:
        if test_type in STYLES:
            styles = STYLES[test_type]
        else:
            styles = STYLES
    step_sum = 0
    step_counts = []
    step_keys = [key for key in steps]
    for step in steps.values():
        step_sum += step
        step_counts.append(step_sum)
    step = -1
    offset = 0
    add_count = 0
    i = 0
    for i in range(len(questions)):
        questions_id = questions_start + i + 1
        if step == -1:
            step = 0
            paragraph.style = styles["title"]
            paragraph.text = step_keys[step]
            paragraph = insert_paragraph_after(paragraph, "", styles['title'])
            offset += 1
        elif i >= step_counts[step]:
            step += 1
            if step >= len(steps):
                print(test_type, "Пропущено вопросов:", len(questions) - i)
                break
            paragraph = insert_paragraph_after(paragraph, step_keys[step], styles['title'])
            paragraph = insert_paragraph_after(paragraph, "", styles['title'])
            offset += 2
        add_count += 1
        question = questions[i]
        if test_type == TYPE_CLOSED_QUESTIONS:
            paragraph = insert_paragraph_after(paragraph, f"{questions_id}\t{question.question_text}",
                                               styles["question"])
            answer_listing = ord("А")
            for answer in question.answers:
                paragraph = insert_paragraph_after(
                    paragraph,
                    f"{chr(answer_listing)})\t{answer}",
                    styles["answer"]
                )
                answer_listing += 1
                offset += 1
            paragraph = insert_paragraph_after(paragraph, "", styles['question'])
            offset += 2
        elif test_type == TYPE_OPENED_QUESTIONS:
            paragraph = insert_paragraph_after(paragraph, f"{questions_id}\t{question.question_text}",
                                               styles["question"])
            paragraph = insert_paragraph_after(paragraph, "", styles['question'])
            offset += 2
        elif test_type == TYPE_ACCORDANCE_QUESTIONS:
            paragraph = insert_paragraph_after(paragraph, f"{questions_id}\tУстановите соответствие",
                                               styles["question"])
            paragraph = insert_paragraph_after(paragraph, f"{question.question_text}", styles["question"])
            table = document.add_table(rows=1, cols=2)
            answer_id = 1
            answer_col1 = 1
            answer_col2 = ord("А")
            row = table.rows[0]
            row.cells[0].text = f"{answer_col1} {question.column1[0]}"
            row.cells[1].text = f"{chr(answer_col2)}) {question.column2[0]}"
            row.cells[0].paragraphs[0].style = styles["table"]
            row.cells[1].paragraphs[0].style = styles["key_answer"]
            for col1, col2 in zip(question.column1[1:], question.column2[1:]):
                row = table.add_row()
                row.cells[0].text = f"{answer_col1 + answer_id} {col1}"
                row.cells[1].text = f"{chr(answer_col2 + answer_id)}) {col2}"
                row.cells[0].paragraphs[0].style = styles["table"]
                row.cells[1].paragraphs[0].style = styles["key_answer"]
                answer_id += 1
            paragraph2 = insert_paragraph_after(paragraph, "", styles['question'])
            move_table_after(table, paragraph)
            paragraph = paragraph2
            offset += 2
    if len(questions) < step_sum:
        print(test_type, "Не хватило вопросов:", step_sum-len(questions))
    return offset, add_count
