import docx
import docxtpl
import regex as re

from parser import *

template_file = "./template/template.docx"
comp = "Какая-то компетенция"
indicator = "Какой-то индикатор"
subject = "Психология"
print_question = False

# Инициализация стилей и шагов
initialize(styles={
    "question": "!СКИФ-ЗТ-Вопрос",
    "answer": "!СКИФ-ЗТ-Ответ",
    "title": "!СКИФ-Категория",
    "table": "!СКИФ-Соответствие",
    "key_num": "!СКИФ-Ключ-Номер",
    "key_answer": "!СКИФ-Ключ-Ответ",
}, steps={
    TYPE_CLOSED_QUESTIONS: {
        "Простые (1 уровень)": 5,
        "Средне–сложные (2 уровень)": 17,
        "Сложные (3 уровень)": 3,
    },
    TYPE_OPENED_QUESTIONS: {
        "Простые (1 уровень)": 2,
        "Средне–сложные (2 уровень)": 7,
        "Сложные (3 уровень)": 1,
    },
    TYPE_ACCORDANCE_QUESTIONS: {
        "Простые (1 уровень)": 7,
        "Средне–сложные (2 уровень)": 24,
        "Сложные (3 уровень)": 4,
    }
})


# Загрузка шаблона
template = docx.Document(template_file)

# Закрытый тест
closed_id = find_paragraph_id(template, "{{ ЗАКРЫТЫЙ ТЕСТ }}", find_types["equal"])
closed_id_end = 0
closed_questions = []
if closed_id >= 0:
    paragraph = template.paragraphs[closed_id]
    closed_questions = load_questions_from_file(
        "./source/closed.txt",
        TYPE_CLOSED_QUESTIONS,
        question_handler=lambda q: re.split('[\t\s\)\.]+', q, 1)[-1].strip(),
        answer_handler=lambda a: re.split('[\t\s\)\.]+', a, 1)[-1].strip(),
        right_answer_handler=lambda a: a.replace("Ответ:", "").strip()
    )
    if print_question:
        print("Закрытый тест:", end="")
        for question in closed_questions:
            print(question, end="\n\n")
    closed_id_end, count = replace_test(template, paragraph, TYPE_CLOSED_QUESTIONS, closed_questions)
    closed_id_end += closed_id
    closed_questions = closed_questions[:count]
else:
    print("Закрытый тест не найден")

# Тест на соответствие
accordance_id = find_paragraph_id(template, "{{ ТЕСТ НА СООТВЕТСТВИЕ }}", find_types["equal"], start_id=closed_id_end)
accordance_id_end = 0
accordance_questions = []
if accordance_id >= 0:
    paragraph = template.paragraphs[accordance_id]
    accordance_questions = load_questions_from_file(
        "./source/Тест_Психология_последняя_версия (2).txt",
        TYPE_ACCORDANCE_QUESTIONS
    )
    if print_question:
        print("Тест на соответствие:", end="")
        for question in accordance_questions:
            print(question, end="\n\n")
    accordance_id_end, count = replace_test(template, paragraph, TYPE_ACCORDANCE_QUESTIONS, accordance_questions,
                                            questions_start=len(closed_questions))
    accordance_questions = accordance_questions[:count]
    accordance_id_end += accordance_id
else:
    print("Тест на соответствие не найден")

# Открытый тест
opened_id = find_paragraph_id(template, "{{ ОТКРЫТЫЙ ТЕСТ }}", find_types["equal"], start_id=accordance_id_end)
opened_id_end = 0
opened_questions = []
if opened_id >= 0:
    paragraph = template.paragraphs[opened_id]
    opened_questions = load_questions_from_file(
        "./source/open.txt",
        TYPE_OPENED_QUESTIONS,
        question_handler=lambda q: re.split('[\t\s\)\.]+', q, 1)[-1].strip(),
        answer_handler=lambda a: re.split('[\t\s\)\.]+', a, 1)[-1].strip(),
        right_answer_handler=lambda a: a[6:].strip() if a.lower().startswith("ответ:") else a.strip()
    )
    if print_question:
        print("Открытый тест:", end="")
        for question in opened_questions:
            print(question, end="\n\n")
    opened_id_end, count = replace_test(template, paragraph, TYPE_OPENED_QUESTIONS, opened_questions,
                                        questions_start=len(closed_questions) + len(accordance_questions))
    opened_questions = opened_questions[:count]
    opened_id_end += opened_id
else:
    print("Открытый тест не найден")

key_table = template.tables[-1]
all_questions = closed_questions + accordance_questions + opened_questions
i = 1
q_i = 0
column_num = 0
column_text = 1
while q_i < len(all_questions):
    row = key_table.rows[i]
    row.cells[column_num].text = str(q_i + 1)
    row.cells[column_num].paragraphs[0].style = STYLES["key_num"]
    row.cells[column_text].text = all_questions[q_i].right()
    row.cells[column_text].paragraphs[0].style = STYLES["key_answer"]
    i += 1
    q_i += 1
    if i >= len(key_table.rows):
        i = 0
        column_num = 4
        column_text = 5

template.save("./tmp/temp.docx")
context = {
    "subject": subject,
    "comp": comp,
    "indicator": indicator,
    "simple_c": STEPS[TYPE_CLOSED_QUESTIONS]["Простые (1 уровень)"],
    "simple_a": STEPS[TYPE_ACCORDANCE_QUESTIONS]["Простые (1 уровень)"],
    "simple_o": STEPS[TYPE_OPENED_QUESTIONS]["Простые (1 уровень)"],
    "medium_c": STEPS[TYPE_CLOSED_QUESTIONS]["Средне–сложные (2 уровень)"],
    "medium_a": STEPS[TYPE_ACCORDANCE_QUESTIONS]["Средне–сложные (2 уровень)"],
    "medium_o": STEPS[TYPE_OPENED_QUESTIONS]["Средне–сложные (2 уровень)"],
    "hard_c": STEPS[TYPE_CLOSED_QUESTIONS]["Сложные (3 уровень)"],
    "hard_a": STEPS[TYPE_ACCORDANCE_QUESTIONS]["Сложные (3 уровень)"],
    "hard_o": STEPS[TYPE_OPENED_QUESTIONS]["Сложные (3 уровень)"],
}
context["all_c"] = context["simple_c"] + context["medium_c"] + context["hard_c"]
context["all_a"] = context["simple_a"] + context["medium_a"] + context["hard_a"]
context["all_o"] = context["simple_o"] + context["medium_o"] + context["hard_o"]
context["all_simple"] = context["simple_c"] + context["simple_a"] + context["simple_o"]
context["all_medium"] = context["medium_c"] + context["medium_a"] + context["medium_o"]
context["all_hard"] = context["hard_c"] + context["hard_a"] + context["hard_o"]
context["all"] = context["all_simple"] + context["all_medium"] + context["all_hard"]
template = docxtpl.DocxTemplate("./tmp/temp.docx")
template.render(context)
template.save(f"./dest/{subject}.docx")
